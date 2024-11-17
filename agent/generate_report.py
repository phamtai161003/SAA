from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import re

def parse_log_file(log_file):
    results = []
    pass_count = 0
    fail_count = 0
    total_criteria = 0
    criteria_name = ""
    command = ""
    output = ""
    error = ""
    
    with open(log_file, 'r') as file:
        for line in file:
            line = line.strip()
            if line.startswith("Command: "):
                command = line.replace("Command: ", "")
            elif line.startswith("Output:"):
                # Đọc phần Output và lưu các dòng cho đến khi gặp dòng trống hoặc Error:
                output = ""
                for output_line in file:
                    if output_line.strip() == "" or output_line.startswith("Error:"):
                        break
                    output += output_line
            elif line.startswith("Error:"):
                error = line.replace("Error: ", "").strip()
                # Xác định kết quả là Pass/Fail dựa trên output và error
                result = "Pass" if "completed successfully" in output or "RUNNING" in output else "Fail"
                # Đếm số Pass và Fail
                if result == "Pass":
                    pass_count += 1
                else:
                    fail_count += 1
                total_criteria += 1
                # Thêm kết quả vào danh sách
                results.append((criteria_name, command, output.strip(), result))
            elif line.startswith("Task: "):
                # Lấy tên tiêu chí từ dòng bắt đầu bằng "Task: "
                criteria_name = line.replace("Task: ", "")
                
    return results, pass_count, fail_count, total_criteria

def generate_report(doc_file, results, pass_count, fail_count, total_criteria):
    doc = Document()
    doc.add_heading("System Compliance Report", 0).alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Phần tổng kết
    summary = doc.add_paragraph()
    summary.add_run("Total Criteria: ").bold = True
    summary.add_run(str(total_criteria))
    summary.add_run("    Pass: ").bold = True
    summary.add_run(str(pass_count))
    summary.add_run("    Fail: ").bold = True
    summary.add_run(str(fail_count))
    doc.add_paragraph("\n")

    # Phần kết quả chi tiết
    doc.add_heading("Detailed Compliance Results", level=1)
    for criteria_name, command, output, result in results:
        doc.add_paragraph(f"Criteria: {criteria_name}", style='Heading2')
        doc.add_paragraph(f"Command: {command}", style='BodyText')
        doc.add_paragraph(f"Output: {output}", style='BodyText')
        doc.add_paragraph(f"Result: {result}", style='BodyText')
        doc.add_paragraph("\n")

    # Lưu tài liệu
    doc.save(doc_file)

# Hàm chính
log_file_path = 'execution_results_windows.log'  # Đường dẫn đến file log
doc_file_path = 'Compliance_Report.docx'  # Đường dẫn file báo cáo

# Phân tích file log và tạo báo cáo
results, pass_count, fail_count, total_criteria = parse_log_file(log_file_path)
generate_report(doc_file_path, results, pass_count, fail_count, total_criteria)
print(f"Report generated: {doc_file_path}")
