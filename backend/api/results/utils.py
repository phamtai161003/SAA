from docx import Document
import re

def generate_report_from_json(data, file_path):
    """
    Tạo báo cáo Word từ dữ liệu JSON.
    """
    document = Document()
    document.add_heading("Báo cáo kiểm tra hệ thống", level=1)

    def add_category_to_report(category, level=1):
        document.add_heading(category["category"], level=level)

        # Add tasks
        for task in category.get("tasks", []):
            document.add_heading(f"- Task: {task['task']}", level=level + 1)
            document.add_paragraph(f"Combine: {task['combine']}")
            document.add_paragraph(f"Scored: {'Yes' if task['scored'] else 'No'}")
            document.add_paragraph(f"Remediation: {task['remediation']}")
            document.add_paragraph(f"Note: {task['note']}")

            # Add commands
            for command in task.get("commands", []):
                document.add_paragraph(f"Command: {command['command']}")
                document.add_paragraph(f"Expected Output: {command['expected_output']}")
                document.add_paragraph(f"Operator: {command['operator']}")
                document.add_paragraph(f"Parser: {command['parser']}")
                document.add_paragraph(f"Output: {command.get('output', 'Not available')}")

                # Evaluate the result
                result = evaluate_command_output(command["output"], command["expected_output"], command["operator"])
                document.add_paragraph(f"Result: {'PASS' if result else 'FAIL'}")

        # Add child categories
        for child in category.get("children", []):
            add_category_to_report(child, level=level + 1)

    add_category_to_report(data)
    document.save(file_path)


def evaluate_command_output(output, expected_output, operator):
    """
    Đánh giá kết quả lệnh dựa trên output, expected_output và operator.
    """
    try:
        if operator == "Equal":
            return output == expected_output
        elif operator == "Contain":
            return expected_output in output
        elif operator == "Notcontain":
            return expected_output not in output
        elif operator == "Includeoneof":
            return any(item in output for item in expected_output.split(","))
        elif operator == "Allin":
            return all(item in output for item in expected_output.split(","))
        elif operator == "Notequal":
            return output != expected_output
        elif operator == "Regex":
            return re.search(expected_output, output) is not None
        elif operator == "GreaterThan":
            return float(output) > float(expected_output)
        elif operator == "GreaterThanOrEqual":
            return float(output) >= float(expected_output)
        elif operator == "LessThan":
            return float(output) < float(expected_output)
        elif operator == "LessThanOrEqual":
            return float(output) <= float(expected_output)
        elif operator == "Between":
            min_val, max_val = map(float, expected_output.split(","))
            return min_val <= float(output) <= max_val
        else:
            return False
    except Exception:
        return False
